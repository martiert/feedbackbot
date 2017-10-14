import json
import logging

import pymongo
import aiohttp
import asyncio
import tempfile
import docx
import ciscosparkapi

from spark import Server

logger = logging.getLogger('Feedback')
logger.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s %(name)s %(levelname)s: %(message)s')
handler = logging.StreamHandler()
handler.setFormatter(formatter)
logger.addHandler(handler)

help_message = '''
 * help - Show this message
 * ask <question> - Ask a question to all of your customers, also fetches the answers for your previous question
    1. This removes the answers from the database. After this operation you alone have the answers to the previous question
 * get answers - Fetch answers for your current question
 * add customer <customer>: <emails> - Add customer, or emails to existing customer.
    1. customer name can not contain ':' characters
    2. customer name can contain spaces
    3. emails are split with spaces
 * remove customer <customer>: <emails> - Remove emails from existing customer entry
 * remove customer <customer>: all - Completely remove customer
 * list customers - List all your customers
 * list emails <customer> - List emails for given customer
 * give customer <receiver> <customer> - Give customer to receiver
'''

admin_help = '''
 * add contact <email> - Add contact person
 * remove contact <email> - Remove contact person, along with his customers
 * add admin <email> - Create admin, or give admin privileges to existing contact person
 * remove admin <email> - Remove admin privileges from contact person
 * steal customer <from> <customer> - Steal customer from contact person
'''



class Feedback:
    def __init__(self, config):
        self._states = {}
        self._next_id = {}
        self._setup_server(config)
        client = pymongo.MongoClient('mongodb://127.0.0.1')
        collection = client[config['database']]
        self._db = collection['contacts']
        self._customers = collection['customers']

    async def answer(self, api, message):
        customer = self._customers.find_one({'_id': message.personEmail})
        if not customer:
            return

        answers = customer.get('answers', [])
        answers.append(message.text)
        self._customers.update({'_id': message.personEmail}, {'$set': {'answers': answers}})

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            'Thank you')

    async def get_answers(self, api, message):
        contact = self._db.find_one({'_id': message.personEmail})
        if not contact:
            answer(api, message)
            return

        logger.info('{} is fetching answers'.format(message.personEmail))
        old_question = contact.get('question', None)
        if old_question:
            await self._send_document(api, old_question, message.personEmail)
            return

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            'You need to ask a question before fetching answers')

    async def ask(self, api, message):
        contact = self._db.find_one({'_id': message.personEmail})
        if not contact:
            answer(api, message)
            return

        old_question = contact.get('question', None)
        if old_question:
            await self._send_document(api, old_question, message.personEmail)

        self._customers.update(
            {'contact': message.personEmail},
            {'$unset': {'answers': 1}},
            multi=True)
        question = ' '.join(message.text.split(' ')[1:]).strip()
        self._db.update({'_id': message.personEmail}, {'$set': {'question': question}})

        logger.info('{} is asking new question'.format(message.personEmail))

        loop = asyncio.get_event_loop()
        for customer in self._customers.find({'contact': message.personEmail}):
            try:
                await loop.run_in_executor(
                    None,
                    api.messages.create,
                    None,
                    None,
                    customer['_id'],
                    question)
            except ciscosparkapi.exceptions.SparkApiError:
                logger.warn('Can not send new question to {}'.format(customer['_id']))
                await loop.run_in_executor(
                    None,
                    api.messages.create,
                    None,
                    None,
                    message.personEmail,
                    'Not able to send messages to {}'.format(customer['_id']))

        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            'All customers asked')

    async def steal_customer(self, api, message):
        loop = asyncio.get_event_loop()
        contact = self._db.find_one({'_id': message.personEmail})
        if not contact:
            answer(api, message)
            return

        data = message.text.replace('steal customer', '').strip().split(' ')
        victim = data[0].strip()
        customer = ' '.join(data[1:]).strip()

        logger.info('{} stole customer {} from {}'.format(message.personEmail, customer, victim))
        self._move_customer(victim, message.personEmail, customer)
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            'Stole customer {} from {}'.format(customer, victim))

    async def give_customer(self, api, message):
        loop = asyncio.get_event_loop()
        contact = self._db.find_one({'_id': message.personEmail})
        if not contact:
            answer(api, message)
            return

        data = message.text.replace('give customer', '').strip().split(' ')
        receiver = data[0].strip()
        customer = ' '.join(data[1:]).strip()

        logger.info('{} gave customer {} to {}'.format(message.personEmail, customer, receiver))
        self._move_customer(message.personEmail, receiver, customer)
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            'Moved customer {} to {}'.format(customer, receiver))

    async def list_customers(self, api, message):
        contact = self._db.find_one({'_id': message.personEmail})
        if not contact:
            answer(api, message)
            return

        result = 'Your customers:'
        customers = set(customer['customer'] for customer in self._customers.find({'contact': message.personEmail}))
        for customer in customers:
            result += '\n * {}'.format(customer)

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            None,
            result)

    async def list_emails(self, api, message):
        loop = asyncio.get_event_loop()
        contact = self._db.find_one({'_id': message.personEmail})
        if not contact:
            answer(api, message)
            return

        customer = message.text.replace('list emails', '').strip()
        result = 'Emails registered for customer {}:'.format(customer)
        for customer in self._customers.find({'customer': customer, 'contact': message.personEmail}):
            result += '\n * {}'.format(customer['_id'])

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            None,
            result)

    async def add_customer(self, api, message):
        loop = asyncio.get_event_loop()
        is_contact = self._db.find_one({'_id': message.personEmail})
        if not is_contact:
            answer(api, message)
            return

        data = message.text.replace('add customer', '').strip().split(':')

        if not len(data) > 1:
            logger.warn('{} had wrong format when creating customer: {}'.format(message.text))
            await loop.run_in_executor(
                None,
                api.messages.create,
                None,
                None,
                message.personEmail,
                'Wrong format. Expected <customer name>: mail1 mail2 mail3')
            return

        customer = data[0]
        emails = set(''.join(data[1:]).strip().split(' '))
        if '' in emails:
            emails.remove('')

        logger.info('{} added {} to customer {}'.format(message.personEmail, content, customer))
        question = is_contact.get('question', None)
        for email in emails:
            try:
                if question:
                    await loop.run_in_executor(
                        None,
                        api.messages.create,
                        None,
                        None,
                        email,
                        question)

                self._customers.insert({'_id': email, 'contact': message.personEmail, 'customer': customer})
            except ciscosparkapi.exceptions.SparkApiError:
                logger.warn('{} Failed to send question to new user {}'.format(message.personEmail, email))
                await loop.run_in_executor(
                    None,
                    api.messages.create,
                    None,
                    None,
                    message.personEmail,
                    'Not adding {} as we are not able to send a question to that address'.format(email))
            except pymongo.errors.DuplicateKeyError:
                logger.warn('{} is trying to re add {}'.format(message.personEmail, email))
                await loop.run_in_executor(
                    None,
                    api.messages.create,
                    None,
                    None,
                    message.personEmail,
                    'Not adding {} as the address is already registers'.format(email))

        customers = [customer['_id'] for customer in self._customers.find({'customer': customer, 'contact': message.personEmail})]
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            '{} are now registered with {}'.format(customer, customers))

    async def remove_customer(self, api, message):
        loop = asyncio.get_event_loop()
        is_contact = self._db.find_one({'_id': message.personEmail})
        if not is_contact:
            answer(api, message)
            return

        data = message.text.replace('remove customer', '').strip().split(':')
        customer = data[0]
        content = ''.join(data[1:]).strip()

        logger.info('{} removed {} from customer {}'.format(message.personEmail, content, customer))
        self._remove_customer(message.personEmail, customer, content)

        customers = [customer['_id'] for customer in self._customers.find({'customer': customer, 'contact': message.personEmail})]
        result = '{} is completely removed'.format(customer)
        if customers:
            result = '{} are now registered with {}'.format(customer, customers)

        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            result)

    async def add_admin(self, api, message):
        is_admin = self._db.find_one({'_id': message.personEmail, 'admin': True})
        if not is_admin:
            answer(api, message)
            return

        email = message.text.replace('add admin', '').strip()
        response = '{} is already an admin'
        entry = self._db.find_one({'_id': email})

        if not entry:
            self._db.insert({'_id': email, 'admin': True})
            response = 'created {} as admin'
            logger.info('{} Created admin {}'.format(message.personEmail, email))
        elif not entry.get('admin', False):
            self._db.update({'_id': email}, {'$set': {'admin': True}})
            response = 'Gave admin privileges to {}'
            logger.info('{} Gave admin privileges to {}'.format(message.personEmail, email))

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            email,
            'You have been granted admin privileges')
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            response.format(email))

    async def remove_admin(self, api, message):
        is_admin = self._db.find_one({'_id': message.personEmail, 'admin': True})
        if not is_admin:
            answer(api, message)
            return

        email = message.text.replace('remove admin', '').strip()
        entry = self._db.find_one({'_id': email, 'admin': True})

        if not entry:
            response = 'Could not find admin {}'.format(email)
        else:
            self._db.update({'_id': email}, {'$set': {'admin': False}})
            response = 'Removed admin privileges from {}'
            logger.info('{} removed admin privileges from {}'.format(message.personEmail, email))

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            email,
            'You have lost admin privileges')
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            response.format(email))

    async def add_contact(self, api, message):
        is_admin = self._db.find_one({'_id': message.personEmail, 'admin': True})
        if not is_admin:
            answer(api, message)
            return

        email = message.text.replace('add contact', '').strip()
        entry = self._db.find_one({'_id': email})
        response = 'Contact {} already exists'

        loop = asyncio.get_event_loop()
        if not entry:
            self._db.insert({'_id': email, 'admin': False})
            response = 'Added {} as contact'
            logger.info('{} added {} as customer contact'.format(message.personEmail, email))
            await loop.run_in_executor(
                None,
                api.messages.create,
                None,
                None,
                email,
                'You have been added as a customer contact. Type help to see the commands you have available')

        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            response.format(email))

    async def remove_contact(self, api, message):
        is_admin = self._db.find_one({'_id': message.personEmail, 'admin': True})
        if not is_admin:
            answer(api, message)
            return

        email = message.text.replace('remove contact', '').strip()
        entry = self._db.find_one({'_id': email})
        response = 'Contact {} does not exist'

        if entry:
            logger.info('{} removed {} as customer contact'.format(message.personEmail, email))
            self._db.remove({'_id': email})
            response = 'Removed {} from contacts'
            await loop.run_in_executor(
                None,
                api.messages.create,
                None,
                None,
                email,
                'You have been removed as a customer contact')

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            None,
            message.personEmail,
            response.format(email))

    async def help(self, api, message):
        is_contact = self._db.find_one({'_id': message.personEmail})
        if not is_contact:
            return

        response = help_message
        if is_contact.get('admin', False):
            response += admin_help

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            None,
            api.messages.create,
            None,
            message.personId,
            None,
            None,
            response)

    def _move_customer(self, from_contact, to_contact, customer):
        self._customers.update(
                {'customer': customer, 'contact': from_contact},
                {'$set': {'contact': to_contact}},
                multi=True)

    def _remove_customer(self, contact, customer, to_remove):
        if to_remove == 'all':
            self._customers.remove({'customer': customer, 'contact': contact})
        else:
            emails = to_remove.split(' ')
            for email in emails:
                self._customers.remove({'_id': email, 'customer': customer, 'contact': contact})

    async def _send_document(self, api, old_question, contact):
        document = docx.Document()
        document.add_heading(old_question, 0)

        respondents = self._customers.find({'contact': contact, 'answers': {'$exists': True}})

        for respondent in respondents:
            document.add_heading('Customer: {}, email: {}'.format(respondent['customer'], respondent['_id']), 3)

            for answer in respondent.get('answers', []):
                document.add_paragraph(answer, style='List Number')

        with tempfile.NamedTemporaryFile(prefix='answers_', suffix='.docx') as fd:
            document.save(fd)
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(
                None,
                api.messages.create,
                None,
                None,
                contact,
                'answers',
                None,
                [fd.name])

    def _setup_server(self, config):
        loop = asyncio.get_event_loop()
        self._server = Server(
            config['bot'],
            loop
        )
        self._server.default_message(self.answer)
        self._server.listen('^help$', self.help)
        self._server.listen('^ask ', self.ask)
        self._server.listen('^get answers$', self.get_answers)
        self._server.listen('^list customers$', self.list_customers)
        self._server.listen('^list emails', self.list_emails)
        self._server.listen('^give customer', self.give_customer)
        self._server.listen('^steal customer', self.steal_customer)
        self._server.listen('^add admin', self.add_admin)
        self._server.listen('^add contact', self.add_contact)
        self._server.listen('^add customer', self.add_customer)
        self._server.listen('^remove admin', self.remove_admin)
        self._server.listen('^remove contact', self.remove_contact)
        self._server.listen('^remove customer', self.remove_customer)

        loop.run_until_complete(self._server.setup())

    def run(self):
        loop = asyncio.get_event_loop()
        print('======== Bot Ready ========')
        try:
            loop.run_forever()
        except KeyboardInterrupt:
            pass
        except:
            print(sys.exc_info())
        finally:
            loop.run_until_complete(self._server.cleanup())
